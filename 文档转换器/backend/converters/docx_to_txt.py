from docx import Document
from .base import BaseConverter
from typing import Dict, Any

class DocxToTxtConverter(BaseConverter):
    """DOCX 到 TXT 转换器（优化版 - 参考 conversion_core）
    
    优化内容：
    1. 添加进度回调
    2. 提取表格内容
    3. 保留段落结构
    4. 更好的文本清理
    """
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['txt']
    
    def convert(self, input_path: str, output_path: str, **options) -> Dict[str, Any]:
        """将 DOCX 转换为 TXT（增强版）"""
        try:
            self.validate_input(input_path)
            self.update_progress(input_path, 10)
            
            # 获取选项
            extract_tables = options.get('extract_tables', True)
            preserve_structure = options.get('preserve_structure', True)
            
            # 读取 DOCX
            doc = Document(input_path)
            self.update_progress(input_path, 30)
            
            full_text = []
            
            # 提取段落
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
                elif preserve_structure:
                    # 保留空行以维持结构
                    full_text.append('')
            
            self.update_progress(input_path, 60)
            
            # 提取表格
            if extract_tables and doc.tables:
                full_text.append('\n' + '='*50 + ' TABLES ' + '='*50)
                
                for table_idx, table in enumerate(doc.tables):
                    full_text.append(f'\n[Table {table_idx + 1}]')
                    
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            full_text.append(row_text)
                    
                    full_text.append('')  # 表格之间空行
            
            self.update_progress(input_path, 80)
            
            # 写入 TXT
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(full_text))
            
            self.update_progress(input_path, 100)
            
            return {
                'success': True,
                'output_path': output_path,
                'size': self.get_output_size(output_path),
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables) if extract_tables else 0
            }
            
        except Exception as e:
            self.cleanup_on_error(output_path)
            raise Exception(f"DOCX to TXT conversion failed: {str(e)}")
