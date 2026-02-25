from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
from .base import BaseConverter
from typing import Dict, Any


class HtmlToDocxConverter(BaseConverter):
    """HTML 到 DOCX 转换器（优化版 - 参考 conversion_core）
    
    优化内容：
    1. 添加进度回调
    2. 支持两种模式：源码模式 / 文本提取模式
    3. 更好的文本清理
    """
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['docx', 'doc']
    
    def convert(self, input_path: str, output_path: str, **options) -> Dict[str, Any]:
        """将 HTML 转换为 DOCX"""
        try:
            self.validate_input(input_path)
            self.update_progress(input_path, 10)
            
            # 获取选项
            mode = options.get('mode', 'text')  # 'source' or 'text'
            
            # 读取 HTML
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            self.update_progress(input_path, 30)
            
            # 创建 Word 文档
            doc = Document()
            
            if mode == 'source':
                # 源码模式：保留完整 HTML 源码
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(html_content)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                # 文本提取模式：解析 HTML 提取文本
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # 移除脚本和样式
                for script in soup(['script', 'style']):
                    script.decompose()
                
                self.update_progress(input_path, 50)
                
                # 提取文本
                text = soup.get_text(separator='\n', strip=True)
                
                # 清理多余空行
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                
                # 添加到文档
                for line in lines:
                    doc.add_paragraph(line)
                
                self.update_progress(input_path, 80)
            
            # 保存文档
            doc.save(output_path)
            self.update_progress(input_path, 100)
            
            return {
                'success': True,
                'output_path': output_path,
                'size': self.get_output_size(output_path),
                'mode': mode
            }
            
        except Exception as e:
            self.cleanup_on_error(output_path)
            raise Exception(f"HTML to DOCX conversion failed: {str(e)}")
