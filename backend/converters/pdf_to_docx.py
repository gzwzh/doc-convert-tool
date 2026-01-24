from pdf2docx import Converter
from .base import BaseConverter
from typing import Dict, Any, List
import os

class PdfToDocxConverter(BaseConverter):
    """PDF 到 DOCX 转换器（优化版 - 参考 conversion_core）
    
    优化内容：
    1. 添加进度回调支持
    2. 添加多策略降级（pdf2docx → PyMuPDF）
    3. 添加页码范围选择
    4. 添加图片提取选项
    5. 更详细的错误处理
    """
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['docx', 'doc']
    
    def convert(self, input_path: str, output_path: str, **options) -> Dict[str, Any]:
        """将 PDF 转换为 DOCX（多策略）"""
        try:
            self.validate_input(input_path)
            self.update_progress(input_path, 5)
            
            # 确保输出目录存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 获取选项
            extract_images = options.get('extract_images', True)
            
            # 获取页码范围 (支持多种参数名)
            raw_page_range = options.get('pdf_page_range') or options.get('page_range')
            print(f"[PdfToDocx] Options received: {options.keys()}")
            print(f"[PdfToDocx] Raw page range: '{raw_page_range}'")
            
            # 解析页码范围 (此时还不知道总页数，先不传 total_pages)
            pages = self.parse_page_range(raw_page_range)
            print(f"[PdfToDocx] Parsed pages: {pages}")
            
            self.update_progress(input_path, 10)
            
            # 策略1: 使用 pdf2docx（推荐，效果最好）
            try:
                result = self._convert_with_pdf2docx(
                    input_path, output_path, 
                    extract_images=extract_images,
                    pages=pages
                )
                self.update_progress(input_path, 100)
                return result
            except Exception as e1:
                print(f"[pdf2docx failed] {e1}, trying PyMuPDF fallback...")
                self.update_progress(input_path, 30)
                
                # 策略2: 使用 PyMuPDF 降级方案
                try:
                    result = self._convert_with_pymupdf(
                        input_path, output_path,
                        extract_images=extract_images,
                        pages=pages
                    )
                    self.update_progress(input_path, 100)
                    result['method'] = 'pymupdf_fallback'
                    result['warning'] = 'pdf2docx failed, used PyMuPDF fallback (basic text extraction)'
                    return result
                except Exception as e2:
                    raise Exception(f"All conversion methods failed. pdf2docx: {str(e1)}, PyMuPDF: {str(e2)}")
            
        except Exception as e:
            self.cleanup_on_error(output_path)
            raise Exception(f"PDF to DOCX conversion failed: {str(e)}")
    
    def _convert_with_pdf2docx(self, input_path: str, output_path: str, 
                               extract_images: bool = True, 
                               pages: List[int] = None) -> Dict[str, Any]:
        """使用 pdf2docx 转换（主策略）"""
        cv = Converter(input_path)
        
        self.update_progress(input_path, 20)
        
        # 转换
        # pdf2docx 的 convert 方法支持 pages 参数 (list of page indices, 0-based)
        if pages:
            print(f"[PdfToDocx] Converting specific pages: {pages}")
            cv.convert(output_path, pages=pages)
        else:
            print(f"[PdfToDocx] Converting all pages")
            cv.convert(output_path)
            
        cv.close()
        
        self.update_progress(input_path, 90)
        
        return {
            'success': True,
            'output_path': output_path,
            'size': self.get_output_size(output_path),
            'method': 'pdf2docx'
        }
    
    def _convert_with_pymupdf(self, input_path: str, output_path: str,
                              extract_images: bool = True,
                              pages: List[int] = None) -> Dict[str, Any]:
        """使用 PyMuPDF 降级方案（基础文本提取）"""
        try:
            import fitz  # PyMuPDF
            from docx import Document
            from docx.shared import Pt, Inches
            from io import BytesIO
        except ImportError as e:
            raise Exception(f"PyMuPDF or python-docx not installed: {e}")
        
        doc_pdf = fitz.open(input_path)
        doc_word = Document()
        
        total_pages = len(doc_pdf)
        
        # 确定要转换的页码列表
        if pages:
            # 过滤有效的页码
            target_pages = [p for p in pages if 0 <= p < total_pages]
            if not target_pages:
                print("[Warning] No valid pages selected, converting all pages")
                target_pages = list(range(total_pages))
        else:
            target_pages = list(range(total_pages))
            
        self.update_progress(input_path, 40)
        
        for i, page_num in enumerate(target_pages):
            page = doc_pdf[page_num]
            
            # 提取文本
            text = page.get_text("text")
            if text.strip():
                doc_word.add_paragraph(text)
            
            # 提取图片（如果启用）
            if extract_images:
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc_pdf.extract_image(xref)
                        if base_image:
                            image_bytes = base_image["image"]
                            img_stream = BytesIO(image_bytes)
                            
                            para = doc_word.add_paragraph()
                            run = para.add_run()
                            run.add_picture(img_stream, width=Inches(4))
                    except Exception as e:
                        print(f"[Image extraction failed] Page {page_num+1}, Image {img_index+1}: {e}")
            
            # 更新进度
            progress = 40 + int(((i + 1) / len(target_pages)) * 50)
            self.update_progress(input_path, progress)
            
            # 添加分页符 (除了最后一页)
            if i < len(target_pages) - 1:
                doc_word.add_page_break()
        
        doc_pdf.close()
        doc_word.save(output_path)
        
        return {
            'success': True,
            'output_path': output_path,
            'size': self.get_output_size(output_path)
        }
