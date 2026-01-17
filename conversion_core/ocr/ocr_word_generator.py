#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
OCR Word生成器 - 基于OCR结果生成Word文档
支持: LaTeX公式渲染、HTML表格、Markdown表格、表格去重
"""

import os
import re
import tempfile
import time
import hashlib
from typing import Dict, List, Optional, Any, Set
from dataclasses import dataclass

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from conversion_core.services.markdown_parser import MarkdownParser, TableData, HeadingData, FormulaData, HTMLTableParser


@dataclass
class OCRPageResult:
    """OCR页面结果模型"""
    page_number: int
    markdown_text: str
    tables: List[Dict]
    images: List[Dict]
    success: bool
    error: Optional[str] = None


class OCRWordGenerator:
    """基于OCR结果生成Word文档"""
    
    HEADING_STYLES = {
        1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3',
        4: 'Heading 4', 5: 'Heading 5', 6: 'Heading 6'
    }
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        self.parser = MarkdownParser()
        self._latex_renderer = None
        self._table_hashes: Set[str] = set()  # 用于表格去重
        self._extracted_images: Dict[int, List[bytes]] = {}  # 缓存提取的图片 {page_num: [image_bytes]}
        
        # 编译正则表达式
        self.HTML_TABLE_PATTERN = re.compile(r'<table[^>]*>.*?</table>', re.DOTALL | re.IGNORECASE)
        self.BLOCK_FORMULA_PATTERN = re.compile(r'\$\$(.+?)\$\$', re.DOTALL)
        self.INLINE_FORMULA_PATTERN = re.compile(r'(?<!\$)\$(?!\$)([^\$]+?)\$(?!\$)')
        # 支持方括号和花括号格式的公式
        self.BRACKET_FORMULA_PATTERN = re.compile(r'\[([^[\]]*(?:\\[^[\]]*)*[^[\]]*)\]')
        self.BRACE_FORMULA_PATTERN = re.compile(r'\{([^{}]*(?:\\[^{}]*)*[^{}]*)\}')
        self.HEADING_PATTERN = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)
        
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
    
    def _get_latex_renderer(self):
        if self._latex_renderer is None:
            self._latex_renderer = LaTeXRenderer()
        return self._latex_renderer
    
    def _get_table_hash(self, headers: List[str], rows: List[List[str]]) -> str:
        """计算表格内容的哈希值，用于去重"""
        # 标准化表格内容
        normalized_headers = [str(h).strip().lower() for h in headers]
        normalized_rows = []
        for row in rows:
            normalized_row = [str(cell).strip().lower() for cell in row]
            normalized_rows.append(normalized_row)
        
        # 创建内容字符串
        content = f"headers:{normalized_headers}|rows:{normalized_rows}"
        return hashlib.md5(content.encode('utf-8')).hexdigest()

    def _process_image_alpha(self, image_bytes: bytes) -> bytes:
        """处理图片alpha通道，将透明背景合成为白色背景
        
        解决印章等带透明背景的图片在Word中显示为黑色背景的问题
        """
        try:
            from PIL import Image
            from io import BytesIO
            
            img = Image.open(BytesIO(image_bytes))
            
            # 检查是否有alpha通道
            if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                # 创建白色背景
                background = Image.new('RGB', img.size, (255, 255, 255))
                
                # 转换为RGBA以获取alpha通道
                if img.mode != 'RGBA':
                    img = img.convert('RGBA')
                
                # 使用alpha通道作为mask合成
                background.paste(img, mask=img.split()[3])
                img = background
            elif img.mode == 'CMYK':
                # CMYK转RGB
                img = img.convert('RGB')
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # 输出为PNG
            output = BytesIO()
            img.save(output, format='PNG')
            return output.getvalue()
            
        except Exception as e:
            print(f"[图片Alpha处理] 处理失败，返回原图: {e}")
            return image_bytes

    def _extract_images_from_pdf(self, pdf_path: str) -> Dict[int, List[bytes]]:
        """从PDF中提取图片
        
        Args:
            pdf_path: PDF文件路径
            
        Returns:
            dict: {page_num: [image_bytes_list]}
        """
        images_by_page = {}
        
        try:
            import fitz  # PyMuPDF
        except ImportError:
            print("[OCR Word生成器] PyMuPDF不可用，无法提取PDF图片")
            return images_by_page
        
        try:
            pdf_doc = fitz.open(pdf_path)
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                image_list = page.get_images(full=True)
                
                if image_list:
                    images_by_page[page_num] = []
                    print(f"[OCR Word生成器] 第{page_num + 1}页发现 {len(image_list)} 张图片")
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = pdf_doc.extract_image(xref)
                            
                            if base_image:
                                image_bytes = base_image["image"]
                                # 处理alpha通道，解决印章背景变黑问题
                                image_bytes = self._process_image_alpha(image_bytes)
                                images_by_page[page_num].append(image_bytes)
                                print(f"[OCR Word生成器] 成功提取第{page_num + 1}页图片 {img_index + 1}")
                        except Exception as e:
                            print(f"[OCR Word生成器] 提取图片失败: {e}")
            
            pdf_doc.close()
            
        except Exception as e:
            print(f"[OCR Word生成器] 提取PDF图片失败: {e}")
        
        return images_by_page
    
    def _extract_seals_from_pdf(self, pdf_path: str, seals_by_page: Dict[int, List[Dict]]) -> Dict[int, List[bytes]]:
        """根据OCR识别的印章区域从PDF裁剪印章图片
        
        Args:
            pdf_path: PDF文件路径
            seals_by_page: {page_num: [seal_dict]} 每页的印章区域信息
            
        Returns:
            dict: {page_num: [seal_image_bytes]}
        """
        seal_images_by_page = {}
        
        if not seals_by_page:
            return seal_images_by_page
        
        try:
            import fitz  # PyMuPDF
        except ImportError:
            print("[OCR Word生成器] PyMuPDF不可用，无法提取印章图片")
            return seal_images_by_page
        
        try:
            pdf_doc = fitz.open(pdf_path)
            
            for page_num, seals in seals_by_page.items():
                if page_num >= len(pdf_doc):
                    continue
                    
                page = pdf_doc[page_num]
                seal_images_by_page[page_num] = []
                
                for seal_idx, seal in enumerate(seals):
                    try:
                        bbox = seal.get('bbox', [])
                        if len(bbox) != 4:
                            continue
                        
                        # bbox格式: [x0, y0, x1, y1]
                        rect = fitz.Rect(bbox)
                        
                        # 使用get_pixmap裁剪印章区域，alpha=False确保白色背景
                        mat = fitz.Matrix(2.0, 2.0)  # 2倍缩放提高质量
                        pix = page.get_pixmap(matrix=mat, clip=rect, alpha=False)
                        
                        # 转换为PNG
                        seal_bytes = pix.tobytes("png")
                        seal_images_by_page[page_num].append(seal_bytes)
                        print(f"[OCR Word生成器] 成功提取第{page_num + 1}页印章 {seal_idx + 1}")
                        
                        pix = None
                    except Exception as e:
                        print(f"[OCR Word生成器] 提取印章失败: {e}")
            
            pdf_doc.close()
            
        except Exception as e:
            print(f"[OCR Word生成器] 提取印章图片失败: {e}")
        
        return seal_images_by_page
    
    def _add_images_to_doc(self, doc: Document, images: List[bytes]):
        """将图片添加到Word文档
        
        Args:
            doc: Word文档对象
            images: 图片字节数据列表
        """
        from io import BytesIO
        
        for img_bytes in images:
            try:
                img_stream = BytesIO(img_bytes)
                
                # 尝试获取图片尺寸
                try:
                    from PIL import Image
                    pil_img = Image.open(BytesIO(img_bytes))
                    width, height = pil_img.size
                    # 计算适当的宽度（最大6英寸）
                    img_width = min(width / 96, 6)  # 假设96 DPI
                except:
                    img_width = 4  # 默认4英寸
                
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(img_stream, width=Inches(img_width))
                
            except Exception as e:
                print(f"[OCR Word生成器] 添加图片到文档失败: {e}")

    def generate(self, ocr_results: List[Any], output_filename: str, target_format: str = 'docx', 
                 source_pdf_path: str = None) -> Dict:
        """生成Word文档
        
        Args:
            ocr_results: OCR识别结果列表
            output_filename: 输出文件名
            target_format: 目标格式 (docx/doc)，OCR模式只支持docx
            source_pdf_path: 源PDF文件路径，用于提取图片
            
        Returns:
            dict: 包含success, output_file, error, warning等字段
        """
        try:
            doc = Document()
            self._set_default_font(doc)
            self._table_hashes.clear()  # 重置表格哈希
            self._extracted_images.clear()  # 清空图片缓存
            
            # OCR模式只支持DOCX格式，如果用户选择DOC格式，给出警告
            warning_msg = None
            if target_format == 'doc':
                warning_msg = "OCR模式仅支持DOCX格式输出，已自动转换为DOCX格式"
                print(f"[OCR Word生成器] {warning_msg}")
            
            # 从源PDF提取图片
            if source_pdf_path and os.path.exists(source_pdf_path):
                self._extracted_images = self._extract_images_from_pdf(source_pdf_path)
            
            # 收集每页的印章区域信息
            seals_by_page = {}
            for idx, page_result in enumerate(ocr_results):
                if self._is_successful_page(page_result):
                    page_num = self._get_page_number(page_result, idx)
                    seals = self._get_seals(page_result)
                    if seals:
                        seals_by_page[page_num] = seals
            
            # 从PDF提取印章图片
            extracted_seals = {}
            if source_pdf_path and os.path.exists(source_pdf_path) and seals_by_page:
                extracted_seals = self._extract_seals_from_pdf(source_pdf_path, seals_by_page)
            
            for idx, page_result in enumerate(ocr_results):
                if not self._is_successful_page(page_result):
                    continue
                
                markdown_text = self._get_markdown_text(page_result)
                if markdown_text:
                    self._process_markdown_text(doc, markdown_text)
                
                # 添加该页的图片
                page_num = self._get_page_number(page_result, idx)
                if page_num in self._extracted_images:
                    self._add_images_to_doc(doc, self._extracted_images[page_num])
                
                # 添加该页的印章图片
                if page_num in extracted_seals:
                    self._add_images_to_doc(doc, extracted_seals[page_num])
                
                if idx < len(ocr_results) - 1:
                    next_page = ocr_results[idx + 1]
                    if self._is_successful_page(next_page):
                        doc.add_page_break()
            
            # 强制使用docx扩展名
            if not output_filename.lower().endswith('.docx'):
                # 移除可能的.doc扩展名
                if output_filename.lower().endswith('.doc'):
                    output_filename = output_filename[:-4] + '.docx'
                else:
                    output_filename += '.docx'
            
            output_path = os.path.join(self.output_dir, output_filename) if self.output_dir else output_filename
            doc.save(output_path)
            
            result = {"success": True, "output_file": output_path, "error": None}
            if warning_msg:
                result["warning"] = warning_msg
            return result
            
        except Exception as e:
            import traceback
            return {"success": False, "output_file": None, "error": f"生成Word文档失败: {str(e)}\n{traceback.format_exc()}"}
    
    def _get_page_number(self, page_result: Any, default_idx: int) -> int:
        """获取页面编号（0-based）"""
        if isinstance(page_result, dict):
            # page_number通常是1-based，转换为0-based
            return page_result.get("page_number", default_idx + 1) - 1
        page_num = getattr(page_result, 'page_number', default_idx + 1)
        return page_num - 1
    
    def _is_successful_page(self, page_result: Any) -> bool:
        if isinstance(page_result, dict):
            return page_result.get("success", False)
        return getattr(page_result, 'success', False)
    
    def _get_markdown_text(self, page_result: Any) -> str:
        if isinstance(page_result, dict):
            return page_result.get("markdown_text", "")
        return getattr(page_result, 'markdown_text', "")
    
    def _get_seals(self, page_result: Any) -> List[Dict]:
        """获取页面的印章区域信息"""
        if isinstance(page_result, dict):
            return page_result.get("seals", [])
        return getattr(page_result, 'seals', [])

    def _set_default_font(self, doc: Document):
        try:
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(11)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except Exception:
            pass

    def _process_markdown_text(self, doc: Document, text: str):
        """直接处理原始Markdown文本"""
        # 第一步：提取HTML表格并替换为占位符
        html_tables = []
        def save_html_table(match):
            html_tables.append(match.group(0))
            return f'[HTML_TABLE_{len(html_tables)-1}]'
        
        text = self.HTML_TABLE_PATTERN.sub(save_html_table, text)
        
        # 第二步：移除其他HTML标签（但保留占位符）
        text = re.sub(r'<(?:div|html|body|center)[^>]*>', '', text, flags=re.IGNORECASE)
        text = re.sub(r'</(?:div|html|body|center)>', '', text, flags=re.IGNORECASE)
        text = re.sub(r'<(?!HTML_TABLE)[^>]+>', '', text)
        
        # 第三步：按行处理
        lines = text.split('\n')
        i = 0
        current_para = []
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # HTML表格占位符
            html_table_match = re.match(r'\[HTML_TABLE_(\d+)\]', stripped)
            if html_table_match:
                if current_para:
                    self._add_text_paragraph(doc, ' '.join(current_para))
                    current_para = []
                
                table_idx = int(html_table_match.group(1))
                if table_idx < len(html_tables):
                    self._add_html_table(doc, html_tables[table_idx])
                i += 1
                continue
            
            # 标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if heading_match:
                if current_para:
                    self._add_text_paragraph(doc, ' '.join(current_para))
                    current_para = []
                
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                self._add_heading(doc, level, heading_text)
                i += 1
                continue
            
            # Markdown表格
            if stripped.startswith('|') and stripped.endswith('|'):
                if current_para:
                    self._add_text_paragraph(doc, ' '.join(current_para))
                    current_para = []
                
                table_lines = [stripped]
                i += 1
                while i < len(lines):
                    next_line = lines[i].strip()
                    if next_line.startswith('|') and next_line.endswith('|'):
                        table_lines.append(next_line)
                        i += 1
                    elif not next_line:
                        i += 1
                        break
                    else:
                        break
                
                self._add_markdown_table(doc, table_lines)
                continue
            
            # 块级公式 $$...$$
            if '$$' in stripped:
                if current_para:
                    self._add_text_paragraph(doc, ' '.join(current_para))
                    current_para = []
                
                formula_text = stripped
                if stripped.count('$$') == 1:
                    i += 1
                    while i < len(lines):
                        formula_text += '\n' + lines[i]
                        if '$$' in lines[i]:
                            i += 1
                            break
                        i += 1
                else:
                    i += 1
                
                match = re.search(r'\$\$(.+?)\$\$', formula_text, re.DOTALL)
                if match:
                    self._add_formula(doc, match.group(1).strip(), display=True)
                continue
            
            # 空行
            if not stripped:
                if current_para:
                    self._add_text_paragraph(doc, ' '.join(current_para))
                    current_para = []
                i += 1
                continue
            
            # 普通文本
            current_para.append(stripped)
            i += 1
        
        if current_para:
            self._add_text_paragraph(doc, ' '.join(current_para))

    def _add_text_paragraph(self, doc: Document, text: str):
        """添加文本段落，处理行内公式"""
        if not text or not text.strip():
            return
        
        text = text.strip()
        
        # 检查是否包含任何类型的公式
        has_dollar_formula = '$' in text
        has_bracket_formula = self.BRACKET_FORMULA_PATTERN.search(text)
        has_brace_formula = self.BRACE_FORMULA_PATTERN.search(text)
        
        if not (has_dollar_formula or has_bracket_formula or has_brace_formula):
            para = doc.add_paragraph()
            self._add_formatted_run(para, text)
            para.paragraph_format.space_after = Pt(6)
            return
        
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        
        # 分割文本和公式
        parts = []
        processed_ranges = []
        
        # 处理$公式$
        for match in self.INLINE_FORMULA_PATTERN.finditer(text):
            if not self._is_in_ranges(match.start(), processed_ranges):
                parts.append((match.start(), match.end(), 'formula', match.group(1)))
                processed_ranges.append((match.start(), match.end()))
        
        # 处理[公式]格式
        for match in self.BRACKET_FORMULA_PATTERN.finditer(text):
            if not self._is_in_ranges(match.start(), processed_ranges) and self._is_likely_formula(match.group(1)):
                parts.append((match.start(), match.end(), 'formula', match.group(1)))
                processed_ranges.append((match.start(), match.end()))
        
        # 处理{公式}格式
        for match in self.BRACE_FORMULA_PATTERN.finditer(text):
            if not self._is_in_ranges(match.start(), processed_ranges) and self._is_likely_formula(match.group(1)):
                parts.append((match.start(), match.end(), 'formula', match.group(1)))
                processed_ranges.append((match.start(), match.end()))
        
        # 按位置排序
        parts.sort(key=lambda x: x[0])
        
        # 构建最终的文本和公式序列
        final_parts = []
        last_end = 0
        
        for start, end, part_type, content in parts:
            if start > last_end:
                final_parts.append(('text', text[last_end:start]))
            final_parts.append((part_type, content))
            last_end = end
        
        if last_end < len(text):
            final_parts.append(('text', text[last_end:]))
        
        # 添加到段落
        for part_type, content in final_parts:
            if part_type == 'text' and content.strip():
                self._add_formatted_run(para, content)
            elif part_type == 'formula':
                self._add_inline_formula_to_para(para, content)
    
    def _is_in_ranges(self, pos: int, ranges: List[tuple]) -> bool:
        """检查位置是否在已处理的范围内"""
        return any(start <= pos < end for start, end in ranges)
    
    def _is_likely_formula(self, text: str) -> bool:
        """判断文本是否可能是LaTeX公式"""
        if not text or len(text) < 2:
            return False
        
        # LaTeX命令特征
        latex_indicators = [
            '\\frac', '\\hat', '\\mathtt', '\\sqrt', '\\sum', '\\int', 
            '\\alpha', '\\beta', '\\gamma', '\\delta', '\\epsilon',
            '\\theta', '\\lambda', '\\mu', '\\pi', '\\sigma', '\\phi',
            '\\psi', '\\omega', '\\Gamma', '\\Delta', '\\Theta', '\\Lambda',
            '\\Pi', '\\Sigma', '\\Phi', '\\Psi', '\\Omega',
            '\\left', '\\right', '\\begin', '\\end', '\\cdot', '\\times',
            '\\div', '\\pm', '\\mp', '\\leq', '\\geq', '\\neq', '\\approx',
            '\\infty', '\\partial', '\\nabla', '\\in', '\\subset', '\\cup',
            '\\cap', '\\emptyset', '\\forall', '\\exists', '\\lim', '\\log',
            '\\ln', '\\sin', '\\cos', '\\tan', '\\sec', '\\csc', '\\cot'
        ]
        
        # 检查是否包含LaTeX命令
        for indicator in latex_indicators:
            if indicator in text:
                return True
        
        # 检查是否包含数学符号组合
        math_patterns = [
            r'\\[a-zA-Z]+',  # LaTeX命令
            r'\^[{]?[^}]*[}]?',  # 上标
            r'_[{]?[^}]*[}]?',   # 下标
        ]
        
        for pattern in math_patterns:
            if re.search(pattern, text):
                return True
        
        return False

    def _add_formatted_run(self, para, text: str):
        """添加格式化文本"""
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = para.add_run(part[1:-1])
                run.font.name = 'Consolas'
            else:
                para.add_run(part)

    def _add_formula(self, doc: Document, latex: str, display: bool = True):
        """添加公式（渲染为图片）"""
        renderer = self._get_latex_renderer()
        image_path = renderer.render_to_image(latex, display=display)
        
        para = doc.add_paragraph()
        if display:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if image_path and os.path.exists(image_path):
            try:
                # 读取图片数据到内存
                with open(image_path, 'rb') as f:
                    image_data = f.read()
                
                # 使用BytesIO插入图片
                from io import BytesIO
                image_stream = BytesIO(image_data)
                
                run = para.add_run()
                width = Inches(5) if display else Inches(2)
                run.add_picture(image_stream, width=width)
                
                # 删除临时文件
                try:
                    os.remove(image_path)
                except:
                    pass
                return
            except Exception as e:
                print(f"[公式图片插入失败] {e}")
        
        # 渲染失败，显示公式文本
        if display:
            para.add_run(f"【公式】{latex}")
        else:
            para.add_run(f"[{latex}]")

    def _add_inline_formula_to_para(self, para, latex: str):
        """在段落中添加行内公式"""
        renderer = self._get_latex_renderer()
        image_path = renderer.render_to_image(latex, display=False)
        
        if image_path and os.path.exists(image_path):
            try:
                with open(image_path, 'rb') as f:
                    image_data = f.read()
                
                from io import BytesIO
                image_stream = BytesIO(image_data)
                
                run = para.add_run()
                run.add_picture(image_stream, height=Pt(14))
                
                try:
                    os.remove(image_path)
                except:
                    pass
                return
            except Exception as e:
                print(f"[行内公式图片插入失败] {e}")
        
        para.add_run(f"[{latex}]")

    def _add_heading(self, doc: Document, level: int, text: str):
        """添加标题"""
        level = max(1, min(6, level))
        try:
            doc.add_heading(text, level=level)
        except Exception:
            para = doc.add_paragraph(text)
            if para.runs:
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(16 - level)

    def _add_html_table(self, doc: Document, html: str):
        """添加HTML表格（带去重）"""
        try:
            parser = HTMLTableParser()
            parser.feed(html)
            
            for table_dict in parser.get_tables():
                headers = table_dict.get('headers', [])
                rows = table_dict.get('rows', [])
                
                # 过滤空表格
                if not rows and not headers:
                    continue
                
                if not headers and rows:
                    headers = rows[0]
                    rows = rows[1:] if len(rows) > 1 else []
                
                # 过滤只有一行且内容为空的表格
                if len(rows) == 0 and len(headers) <= 1:
                    continue
                
                # 检查是否重复
                table_hash = self._get_table_hash(headers, rows)
                if table_hash in self._table_hashes:
                    print(f"[跳过重复HTML表格] hash={table_hash[:8]}, headers={len(headers)}, rows={len(rows)}")
                    continue
                
                self._table_hashes.add(table_hash)
                print(f"[添加HTML表格] hash={table_hash[:8]}, headers={len(headers)}, rows={len(rows)}")
                self._create_word_table(doc, headers, rows)
        except Exception as e:
            print(f"[HTML表格解析失败] {e}")

    def _add_markdown_table(self, doc: Document, lines: List[str]):
        """添加Markdown表格（带去重）"""
        rows = []
        headers = []
        separator_found = False
        
        for idx, line in enumerate(lines):
            line = line.strip()
            if line.startswith('|'):
                line = line[1:]
            if line.endswith('|'):
                line = line[:-1]
            
            cells = [c.strip() for c in line.split('|')]
            
            if all(self._is_separator(c) for c in cells):
                separator_found = True
                continue
            
            if not separator_found and idx == 0:
                headers = cells
            else:
                rows.append(cells)
        
        if not separator_found and headers:
            rows.insert(0, headers)
            headers = []
        
        # 过滤空表格
        if not rows and not headers:
            return
        
        # 过滤只有一行且内容为空的表格
        if len(rows) == 0 and len(headers) <= 1:
            return
        
        # 检查是否重复
        table_hash = self._get_table_hash(headers, rows)
        if table_hash in self._table_hashes:
            print(f"[跳过重复Markdown表格] hash={table_hash[:8]}, headers={len(headers)}, rows={len(rows)}")
            return
        
        self._table_hashes.add(table_hash)
        print(f"[添加Markdown表格] hash={table_hash[:8]}, headers={len(headers)}, rows={len(rows)}")
        self._create_word_table(doc, headers, rows)
    
    def _is_separator(self, cell: str) -> bool:
        cell = cell.strip()
        if not cell:
            return True
        return all(c in '-: ' for c in cell) and '-' in cell

    def _create_word_table(self, doc: Document, headers: List[str], rows: List[List[str]]):
        """创建Word表格"""
        all_rows = ([headers] if headers else []) + rows
        if not all_rows:
            return
        
        row_count = len(all_rows)
        col_count = max(len(r) for r in all_rows) if all_rows else 0
        
        if row_count == 0 or col_count == 0:
            return
        
        table = doc.add_table(rows=row_count, cols=col_count)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row_data in enumerate(all_rows):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < col_count:
                    cell = row.cells[col_idx]
                    cell.text = str(cell_text) if cell_text else ""
                    
                    if row_idx == 0 and headers:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True
        
        doc.add_paragraph()


class LaTeXRenderer:
    """LaTeX公式渲染器"""
    
    def __init__(self):
        self._matplotlib_available = None
    
    def _check_matplotlib(self) -> bool:
        if self._matplotlib_available is None:
            try:
                import matplotlib
                matplotlib.use('Agg')
                self._matplotlib_available = True
            except ImportError:
                self._matplotlib_available = False
        return self._matplotlib_available
    
    def render_to_image(self, latex: str, display: bool = True, dpi: int = 150) -> Optional[str]:
        """将LaTeX公式渲染为图片"""
        if not self._check_matplotlib():
            return None
        
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            
            # 清理和预处理LaTeX代码
            latex = self._preprocess_latex(latex)
            
            fig = plt.figure(figsize=(0.01, 0.01))
            fig.patch.set_facecolor('white')
            
            # 根据是否为显示模式选择字体大小
            fontsize = 14 if display else 12
            latex_str = f"${latex}$"
            
            try:
                text = fig.text(0, 0, latex_str, fontsize=fontsize,
                               verticalalignment='bottom',
                               horizontalalignment='left')
            except Exception as e:
                print(f"[LaTeX语法错误] {latex[:50]}... 错误: {e}")
                plt.close(fig)
                return None
            
            fig.canvas.draw()
            renderer = fig.canvas.get_renderer()
            bbox = text.get_window_extent(renderer)
            
            width = bbox.width / dpi + 0.2
            height = bbox.height / dpi + 0.2
            fig.set_size_inches(max(width, 0.5), max(height, 0.3))
            
            text.set_position((0.1, 0.1))
            
            # 创建临时文件
            fd, temp_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            
            fig.savefig(temp_path, dpi=dpi, bbox_inches='tight',
                       pad_inches=0.1, facecolor='white', transparent=False)
            plt.close(fig)
            
            return temp_path
            
        except Exception as e:
            print(f"[LaTeX渲染失败] {latex[:30]}... 错误: {e}")
            return None
    
    def _preprocess_latex(self, latex: str) -> str:
        """预处理LaTeX代码，修复常见问题"""
        if not latex:
            return latex
        
        # 移除多余的空格
        latex = latex.strip()
        
        # 修复常见的OCR识别错误
        replacements = {
            # 修复花括号问题
            '}{': '}{',
            # 修复常见的命令识别错误
            '\\mathtt{e}': 'e',
            '\\hat{c}': '\\hat{c}',
            # 修复分数格式
            '\\frac{1}{\\mathtt{e}}': '\\frac{1}{e}',
            # 修复其他常见问题
            '\\\\': '\\',
        }
        
        for old, new in replacements.items():
            latex = latex.replace(old, new)
        
        # 确保花括号匹配
        open_braces = latex.count('{')
        close_braces = latex.count('}')
        if open_braces > close_braces:
            latex += '}' * (open_braces - close_braces)
        elif close_braces > open_braces:
            latex = '{' * (close_braces - open_braces) + latex
        
        return latex
